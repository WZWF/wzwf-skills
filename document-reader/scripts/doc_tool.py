#!/usr/bin/env python3
"""Document Tool - Word/Excel 文档操作命令行工具
用于 AI Agent 读取和解析 Word (.docx) 和 Excel (.xlsx) 文件。
"""

import argparse
import json
import os
import re
import sys
import base64
import zipfile
from pathlib import Path
from xml.etree import ElementTree
from io import BytesIO

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from openpyxl import load_workbook

NAMESPACES = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'c': 'http://schemas.openxmlformats.org/drawingml/2006/chart',
}


def output(text):
    print(text)


def output_json(data):
    print(json.dumps(data, ensure_ascii=False, indent=2, default=str))


# ==================== Word 核心函数 ====================

def table_to_markdown(table: Table) -> str:
    if not table.rows:
        return ""
    lines = []
    for row_idx, row in enumerate(table.rows):
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        lines.append("| " + " | ".join(cells) + " |")
        if row_idx == 0:
            lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return "\n".join(lines)


def get_relationship_targets(file_path: str) -> dict:
    rels = {}
    with zipfile.ZipFile(file_path, 'r') as z:
        try:
            rels_xml = z.read('word/_rels/document.xml.rels').decode('utf-8')
            root = ElementTree.fromstring(rels_xml)
            for rel in root.findall('.//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
                r_id = rel.get('Id')
                target = rel.get('Target')
                if rel.get('TargetMode', '') == 'External':
                    rels[r_id] = target
        except Exception:
            pass
    return rels


def extract_hyperlinks(paragraph: Paragraph, rels: dict) -> list[dict]:
    hyperlinks = []
    p_xml = paragraph._element
    for hyperlink in p_xml.findall('.//w:hyperlink', NAMESPACES):
        r_id = hyperlink.get(qn('r:id'))
        texts = []
        for t in hyperlink.findall('.//w:t', NAMESPACES):
            if t.text:
                texts.append(t.text)
        link_text = ''.join(texts)
        url = rels.get(r_id, '') if r_id else ''
        if link_text or url:
            hyperlinks.append({"text": link_text, "url": url})
    return hyperlinks


def get_list_info(paragraph: Paragraph) -> dict | None:
    p_xml = paragraph._element
    num_pr = p_xml.find('.//w:numPr', NAMESPACES)
    if num_pr is None:
        return None
    ilvl = num_pr.find('w:ilvl', NAMESPACES)
    level = int(ilvl.get(qn('w:val'))) if ilvl is not None else 0
    num_id_elem = num_pr.find('w:numId', NAMESPACES)
    num_id = int(num_id_elem.get(qn('w:val'))) if num_id_elem is not None else 0
    return {"level": level, "num_id": num_id}


def extract_images_from_docx(file_path: str, save_dir: str | None = None) -> list[dict]:
    images = []
    mime_types = {
        '.png': 'image/png', '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
        '.gif': 'image/gif', '.bmp': 'image/bmp', '.tiff': 'image/tiff',
        '.emf': 'image/emf', '.wmf': 'image/wmf',
    }
    with zipfile.ZipFile(file_path, 'r') as z:
        for fi in z.filelist:
            if fi.filename.startswith('word/media/'):
                name = os.path.basename(fi.filename)
                data = z.read(fi.filename)
                ext = os.path.splitext(name)[1].lower()
                info = {
                    "name": name, "path": fi.filename,
                    "size": len(data), "size_kb": round(len(data) / 1024, 1),
                    "mime_type": mime_types.get(ext, 'application/octet-stream'),
                }
                if save_dir:
                    os.makedirs(save_dir, exist_ok=True)
                    save_path = os.path.join(save_dir, name)
                    with open(save_path, 'wb') as f:
                        f.write(data)
                    info["saved_path"] = save_path
                images.append(info)
    return images


def extract_charts_info(file_path: str) -> list[dict]:
    charts = []
    with zipfile.ZipFile(file_path, 'r') as z:
        chart_files = [f for f in z.namelist() if f.startswith('word/charts/chart') and f.endswith('.xml')]
        for cf in chart_files:
            try:
                xml_str = z.read(cf).decode('utf-8')
                root = ElementTree.fromstring(xml_str)
                info = {"file": cf, "type": "unknown", "title": ""}
                for elem in root.iter():
                    if 'title' in elem.tag.lower():
                        for t in elem.iter():
                            if t.text and t.text.strip():
                                info["title"] = t.text.strip()
                                break
                for ct in ['barChart', 'lineChart', 'pieChart', 'areaChart', 'scatterChart', 'doughnutChart', 'radarChart']:
                    if ct in xml_str:
                        info["type"] = ct.replace('Chart', '')
                        break
                charts.append(info)
            except Exception as e:
                charts.append({"file": cf, "error": str(e)})
    return charts


def extract_docx_content(file_path, include_tables=True, include_images=True,
                         include_hyperlinks=True, include_charts=True, image_save_dir=None):
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")
    if not file_path.lower().endswith('.docx'):
        raise ValueError(f"不支持的文件格式，请提供 .docx 文件: {file_path}")

    doc = Document(file_path)
    rels = get_relationship_targets(file_path) if include_hyperlinks else {}

    result = {
        "file_path": file_path, "file_name": os.path.basename(file_path),
        "paragraphs": [], "tables": [], "images": [],
        "hyperlinks": [], "charts": [], "full_text": "", "metadata": {},
    }

    props = doc.core_properties
    result["metadata"] = {
        "title": props.title or "", "author": props.author or "",
        "subject": props.subject or "",
        "created": str(props.created) if props.created else "",
        "modified": str(props.modified) if props.modified else "",
        "last_modified_by": props.last_modified_by or "",
    }

    text_parts = []
    all_links = []

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        info = {
            "index": idx, "text": text,
            "style": para.style.name if para.style else "Normal",
            "is_list": False, "list_level": 0, "hyperlinks": [],
        }

        list_info = get_list_info(para)
        if list_info:
            info["is_list"] = True
            info["list_level"] = list_info["level"]

        if include_hyperlinks:
            links = extract_hyperlinks(para, rels)
            info["hyperlinks"] = links
            for link in links:
                all_links.append({"paragraph_index": idx, **link})

        result["paragraphs"].append(info)
        if text:
            if info["is_list"]:
                text_parts.append(f"{'  ' * info['list_level']}* {text}")
            else:
                text_parts.append(text)

    result["full_text"] = "\n\n".join(text_parts)
    result["hyperlinks"] = all_links

    if include_tables:
        for ti, table in enumerate(doc.tables):
            td = {"index": ti, "markdown": table_to_markdown(table), "rows": []}
            for row in table.rows:
                td["rows"].append([cell.text.strip() for cell in row.cells])
            result["tables"].append(td)

    if include_images:
        result["images"] = extract_images_from_docx(file_path, image_save_dir)
    if include_charts:
        result["charts"] = extract_charts_info(file_path)

    return result


def format_docx_output(content, fmt="markdown"):
    if fmt == "json":
        content_copy = content.copy()
        if content_copy.get("images"):
            content_copy["images"] = [{k: v for k, v in img.items() if k != "base64"} for img in content_copy["images"]]
        return json.dumps(content_copy, ensure_ascii=False, indent=2, default=str)

    if fmt == "text":
        parts = [f"文件: {content['file_name']}\n", "=" * 50, content["full_text"]]
        if content["tables"]:
            parts.append("\n" + "=" * 50 + "\n表格内容:")
            for t in content["tables"]:
                parts.append(f"\n表格 {t['index'] + 1}:")
                for row in t["rows"]:
                    parts.append(" | ".join(row))
        if content["hyperlinks"]:
            parts.append("\n" + "=" * 50 + "\n超链接:")
            for link in content["hyperlinks"]:
                parts.append(f"  [{link['text']}]({link['url']})")
        if content["images"]:
            parts.append(f"\n图片: 共 {len(content['images'])} 张")
        if content["charts"]:
            parts.append(f"\n图表: 共 {len(content['charts'])} 个")
        return "\n".join(parts)

    # markdown
    parts = [f"# {content['file_name']}\n"]
    if any(content["metadata"].values()):
        parts.append("## 文档信息\n")
        for k, v in content["metadata"].items():
            if v:
                parts.append(f"- **{k}**: {v}")
        parts.append("")

    parts.append("## 文档内容\n")
    parts.append(content["full_text"])

    if content["tables"]:
        parts.append("\n## 表格\n")
        for t in content["tables"]:
            parts.append(f"### 表格 {t['index'] + 1}\n")
            parts.append(t["markdown"])
            parts.append("")

    if content["hyperlinks"]:
        parts.append("\n## 超链接\n")
        for link in content["hyperlinks"]:
            if link["url"]:
                parts.append(f"- [{link['text'] or link['url']}]({link['url']})")
            else:
                parts.append(f"- {link['text']} (内部链接)")
        parts.append("")

    if content["images"]:
        parts.append(f"\n## 图片\n\n共 **{len(content['images'])}** 张:\n")
        for img in content["images"]:
            parts.append(f"- **{img['name']}** ({img['size_kb']} KB, {img['mime_type']})")
            if img.get('saved_path'):
                parts.append(f"  - 已保存: `{img['saved_path']}`")
        parts.append("")

    if content["charts"]:
        parts.append(f"\n## 图表\n\n共 **{len(content['charts'])}** 个:\n")
        for i, c in enumerate(content["charts"], 1):
            parts.append(f"- **图表 {i}**: {c.get('title', '未命名')} (类型: {c.get('type', 'unknown')})")
        parts.append("")

    return "\n".join(parts)


# ==================== Excel 核心函数 ====================

def excel_table_to_markdown(data, has_header=True):
    if not data or not data[0]:
        return ""
    lines = []
    for i, row in enumerate(data):
        cells = [str(c).replace('|', '\\|') if c is not None else "" for c in row]
        lines.append("| " + " | ".join(cells) + " |")
        if i == 0 and has_header:
            lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return "\n".join(lines)


def get_merged_cell_value(ws, row, col, merged_ranges):
    for mr in merged_ranges:
        if mr.min_row <= row <= mr.max_row and mr.min_col <= col <= mr.max_col:
            return ws.cell(row=mr.min_row, column=mr.min_col).value
    return ws.cell(row=row, column=col).value


def read_excel_file(file_path, sheet_name=None, header_row=0, include_empty_rows=False):
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")
    if not file_path.lower().endswith('.xlsx'):
        raise ValueError(f"不支持的文件格式，请提供 .xlsx 文件: {file_path}")

    result = {
        "file_path": file_path, "file_name": os.path.basename(file_path),
        "sheets": [], "metadata": {},
    }

    wb = None
    try:
        wb = load_workbook(file_path, data_only=True, read_only=False)
        result["metadata"] = {"total_sheets": len(wb.sheetnames), "sheet_names": wb.sheetnames}

        sheets_to_read = [sheet_name] if sheet_name else wb.sheetnames

        for sn in sheets_to_read:
            if sn not in wb.sheetnames:
                raise ValueError(f"工作表 '{sn}' 不存在。可用: {wb.sheetnames}")

            ws = wb[sn]
            merged_ranges = list(ws.merged_cells.ranges)

            sheet_data = []
            max_row = ws.max_row or 0
            max_col = ws.max_column or 0

            if max_row > 0 and max_col > 0:
                for ri in range(1, max_row + 1):
                    row_data = []
                    for ci in range(1, max_col + 1):
                        val = get_merged_cell_value(ws, ri, ci, merged_ranges)
                        row_data.append(val if val is not None else "")
                    if not include_empty_rows and all(
                        c == "" or (isinstance(c, str) and c.strip() == "") for c in row_data
                    ):
                        continue
                    sheet_data.append(row_data)

            headers = None
            data_rows = sheet_data
            if header_row is not None and len(sheet_data) > header_row:
                headers = sheet_data[header_row]
                data_rows = sheet_data[header_row + 1:] if header_row < len(sheet_data) - 1 else []

            normalized_rows = data_rows
            if headers:
                max_cols = max(len(headers), max((len(r) for r in data_rows), default=0))
                headers = list(headers) + [f"列{i+1}" for i in range(len(headers), max_cols)]
                normalized_rows = []
                for row in data_rows:
                    nr = list(row) + [""] * (max_cols - len(row))
                    normalized_rows.append(nr[:max_cols])

            sheet_info = {
                "name": sn, "index": wb.sheetnames.index(sn),
                "max_row": max_row, "max_column": max_col,
                "has_header": headers is not None,
                "headers": headers,
                "data": normalized_rows,
                "row_count": len(normalized_rows),
                "column_count": max_col,
                "header_row": header_row if headers else None,
                "markdown": excel_table_to_markdown(
                    [headers] + normalized_rows if headers else normalized_rows,
                    has_header=headers is not None,
                ),
            }
            result["sheets"].append(sheet_info)
    except Exception as e:
        raise Exception(f"读取 Excel 文件时发生错误: {str(e)}")
    finally:
        if wb:
            wb.close()

    return result


def format_excel_output(excel_data, fmt="markdown"):
    if fmt == "json":
        return json.dumps(excel_data, ensure_ascii=False, indent=2, default=str)

    if fmt == "text":
        parts = [f"文件: {excel_data['file_name']}\n", "=" * 50]
        parts.append(f"总工作表数: {excel_data['metadata']['total_sheets']}")
        parts.append(f"工作表: {', '.join(excel_data['metadata']['sheet_names'])}\n")
        for sheet in excel_data["sheets"]:
            parts.append(f"\n工作表: {sheet['name']}")
            parts.append("-" * 50)
            for row in sheet["data"]:
                parts.append(" | ".join(str(c) for c in row))
        return "\n".join(parts)

    # markdown
    parts = [f"# {excel_data['file_name']}\n"]
    parts.append("## 文件信息\n")
    parts.append(f"- 总工作表数: {excel_data['metadata']['total_sheets']}")
    parts.append(f"- 工作表: {', '.join(excel_data['metadata']['sheet_names'])}\n")

    for sheet in excel_data["sheets"]:
        parts.append(f"## 工作表: {sheet['name']}\n")
        parts.append(f"- 行数: {sheet['row_count']}")
        parts.append(f"- 列数: {sheet['column_count']}")
        parts.append(f"- 表头: {'有' if sheet['has_header'] else '无'}\n")
        parts.append(sheet["markdown"])
        parts.append("")

    return "\n".join(parts)


# ==================== 命令处理 ====================

def cmd_read_docx(args):
    fmt = getattr(args, 'format', 'markdown')
    content = extract_docx_content(
        args.file, include_tables=True, include_images=True,
        include_hyperlinks=True, include_charts=True,
        image_save_dir=getattr(args, 'image_save_dir', None),
    )
    output(format_docx_output(content, fmt))


def cmd_docx_metadata(args):
    content = extract_docx_content(
        args.file, include_tables=False, include_images=False,
        include_hyperlinks=False, include_charts=False,
    )
    meta = content["metadata"]
    lines = [f"# 文档元数据: {content['file_name']}\n"]
    for k, v in meta.items():
        if v:
            lines.append(f"- **{k}**: {v}")
    if not any(meta.values()):
        lines.append("_此文档没有元数据信息_")
    output("\n".join(lines))


def cmd_docx_tables(args):
    content = extract_docx_content(
        args.file, include_tables=True, include_images=False,
        include_hyperlinks=False, include_charts=False,
    )
    tables = content["tables"]
    if not tables:
        output("此文档中没有表格")
        return

    lines = [f"# 表格内容: {content['file_name']}\n"]
    idx = getattr(args, 'index', None)
    if idx is not None:
        if 0 <= idx < len(tables):
            lines.append(f"## 表格 {idx + 1}\n")
            lines.append(tables[idx]["markdown"])
        else:
            output(f"表格索引超出范围。共 {len(tables)} 个表格（索引 0-{len(tables)-1}）")
            return
    else:
        for t in tables:
            lines.append(f"## 表格 {t['index'] + 1}\n")
            lines.append(t["markdown"])
            lines.append("")
    output("\n".join(lines))


def cmd_docx_images(args):
    save_dir = getattr(args, 'save_dir', None)
    images = extract_images_from_docx(args.file, save_dir)
    if not images:
        output("此文档中没有图片")
        return

    lines = [f"# 图片列表\n\n共 **{len(images)}** 张:\n"]
    for i, img in enumerate(images, 1):
        lines.append(f"## 图片 {i}: {img['name']}")
        lines.append(f"- 大小: {img['size_kb']} KB")
        lines.append(f"- 类型: {img['mime_type']}")
        if img.get('saved_path'):
            lines.append(f"- 已保存: `{img['saved_path']}`")
        lines.append("")
    output("\n".join(lines))


def cmd_docx_links(args):
    content = extract_docx_content(
        args.file, include_tables=False, include_images=False,
        include_hyperlinks=True, include_charts=False,
    )
    links = content["hyperlinks"]
    if not links:
        output("此文档中没有超链接")
        return

    lines = [f"# 超链接列表: {content['file_name']}\n\n共 **{len(links)}** 个:\n"]
    for i, link in enumerate(links, 1):
        lines.append(f"## 链接 {i}")
        lines.append(f"- 文本: {link['text']}")
        lines.append(f"- URL: {link['url'] or '(内部链接)'}")
        lines.append(f"- 位置: 段落 {link['paragraph_index'] + 1}")
        lines.append("")
    output("\n".join(lines))


def cmd_docx_charts(args):
    charts = extract_charts_info(args.file)
    if not charts:
        output("此文档中没有图表")
        return

    lines = [f"# 图表列表\n\n共 **{len(charts)}** 个:\n"]
    for i, c in enumerate(charts, 1):
        lines.append(f"## 图表 {i}")
        lines.append(f"- 标题: {c.get('title', '未命名')}")
        lines.append(f"- 类型: {c.get('type', 'unknown')}")
        if c.get('error'):
            lines.append(f"- 解析错误: {c['error']}")
        lines.append("")
    output("\n".join(lines))


def cmd_search_docx(args):
    content = extract_docx_content(
        args.file, include_tables=True, include_images=False,
        include_hyperlinks=False, include_charts=False,
    )
    case_sensitive = getattr(args, 'case_sensitive', False)
    search_key = args.query if case_sensitive else args.query.lower()
    results = []

    for para in content["paragraphs"]:
        text = para["text"]
        compare = text if case_sensitive else text.lower()
        if search_key in compare:
            results.append({"type": "paragraph", "index": para["index"], "text": text, "style": para["style"]})

    for table in content["tables"]:
        for ri, row in enumerate(table["rows"]):
            for ci, cell in enumerate(row):
                compare = cell if case_sensitive else cell.lower()
                if search_key in compare:
                    results.append({"type": "table_cell", "table_index": table["index"], "row": ri, "column": ci, "text": cell})

    if not results:
        output(f'未找到包含 "{args.query}" 的内容')
        return

    lines = [f'# 搜索结果: "{args.query}"\n\n在 {content["file_name"]} 中找到 **{len(results)}** 处匹配\n']
    for i, r in enumerate(results, 1):
        if r["type"] == "paragraph":
            lines.append(f"## 匹配 {i}: 段落 {r['index'] + 1}")
            lines.append(f"样式: {r['style']}")
            lines.append(f"> {r['text']}\n")
        else:
            lines.append(f"## 匹配 {i}: 表格 {r['table_index'] + 1}")
            lines.append(f"位置: 第 {r['row'] + 1} 行, 第 {r['column'] + 1} 列")
            lines.append(f"> {r['text']}\n")
    output("\n".join(lines))


def cmd_read_excel(args):
    fmt = getattr(args, 'format', 'markdown')
    has_header = not getattr(args, 'no_header', False)
    header_row = getattr(args, 'header_row', 0) if has_header else None
    sheet = getattr(args, 'sheet', None)
    include_empty = getattr(args, 'include_empty', False)

    data = read_excel_file(args.file, sheet_name=sheet, header_row=header_row, include_empty_rows=include_empty)
    output(format_excel_output(data, fmt))


def cmd_excel_sheets(args):
    data = read_excel_file(args.file)
    lines = [f"# Excel 工作表列表: {data['file_name']}\n\n共 **{len(data['sheets'])}** 个:\n"]
    for i, s in enumerate(data["sheets"], 1):
        lines.append(f"## 工作表 {i}: {s['name']}")
        lines.append(f"- 索引: {s['index']}")
        lines.append(f"- 行数: {s['row_count']}")
        lines.append(f"- 列数: {s['column_count']}")
        lines.append(f"- 表头: {'有' if s['has_header'] else '无'}")
        if s['headers']:
            lines.append(f"- 表头列: {', '.join(str(h) for h in s['headers'])}")
        lines.append("")
    output("\n".join(lines))


def cmd_excel_table(args):
    has_header = not getattr(args, 'no_header', False)
    header_row = getattr(args, 'header_row', 0) if has_header else None
    sheet = getattr(args, 'sheet', None)

    data = read_excel_file(args.file, sheet_name=sheet, header_row=header_row)
    if not data["sheets"]:
        output("未找到工作表")
        return

    s = data["sheets"][0]
    lines = [f"# Excel 表格: {s['name']}\n"]
    lines.append(f"文件: {data['file_name']}\n")
    lines.append(f"- 行数: {s['row_count']}")
    lines.append(f"- 列数: {s['column_count']}\n")
    lines.append(s["markdown"])
    output("\n".join(lines))


def cmd_search_excel(args):
    case_sensitive = getattr(args, 'case_sensitive', False)
    sheet = getattr(args, 'sheet', None)

    data = read_excel_file(args.file, sheet_name=sheet, include_empty_rows=True)
    search_key = args.query if case_sensitive else args.query.lower()
    results = []

    for s in data["sheets"]:
        hr_idx = s.get("header_row", 0) if s["has_header"] else None
        data_start = (hr_idx + 2) if hr_idx is not None else 1

        for ri, row in enumerate(s["data"]):
            for ci, cell in enumerate(row):
                cell_str = str(cell) if cell is not None else ""
                compare = cell_str if case_sensitive else cell_str.lower()
                if search_key in compare:
                    results.append({
                        "sheet": s["name"],
                        "row": data_start + ri,
                        "column": ci + 1,
                        "value": cell_str,
                    })

    if not results:
        output(f'未找到包含 "{args.query}" 的内容')
        return

    lines = [f'# 搜索结果: "{args.query}"\n\n在 {data["file_name"]} 中找到 **{len(results)}** 处匹配\n']

    by_sheet = {}
    for r in results:
        by_sheet.setdefault(r["sheet"], []).append(r)

    for sn, sresults in by_sheet.items():
        lines.append(f"## 工作表: {sn}")
        for i, r in enumerate(sresults, 1):
            lines.append(f"### 匹配 {i}")
            lines.append(f"- 位置: 第 {r['row']} 行, 第 {r['column']} 列")
            lines.append(f"- 值: {r['value']}\n")

    output("\n".join(lines))


# ==================== CLI 入口 ====================

def main():
    parser = argparse.ArgumentParser(description='Document Tool - Word/Excel 文档操作命令行工具')
    subparsers = parser.add_subparsers(dest='command', help='可用命令')

    # Word 命令
    p = subparsers.add_parser('read-docx', help='读取 Word 文档完整内容')
    p.add_argument('-f', '--file', required=True, help='docx 文件路径')
    p.add_argument('--format', choices=['markdown', 'text', 'json'], default='markdown')
    p.add_argument('--image-save-dir', help='图片保存目录')

    p = subparsers.add_parser('docx-metadata', help='获取文档元数据')
    p.add_argument('-f', '--file', required=True, help='docx 文件路径')

    p = subparsers.add_parser('docx-tables', help='提取文档表格')
    p.add_argument('-f', '--file', required=True, help='docx 文件路径')
    p.add_argument('--index', type=int, help='指定表格索引（从0开始）')

    p = subparsers.add_parser('docx-images', help='提取文档图片')
    p.add_argument('-f', '--file', required=True, help='docx 文件路径')
    p.add_argument('--save-dir', help='图片保存目录')

    p = subparsers.add_parser('docx-links', help='提取文档超链接')
    p.add_argument('-f', '--file', required=True, help='docx 文件路径')

    p = subparsers.add_parser('docx-charts', help='提取文档图表')
    p.add_argument('-f', '--file', required=True, help='docx 文件路径')

    p = subparsers.add_parser('search-docx', help='搜索 Word 文档')
    p.add_argument('-f', '--file', required=True, help='docx 文件路径')
    p.add_argument('-q', '--query', required=True, help='搜索关键词')
    p.add_argument('--case-sensitive', action='store_true', help='区分大小写')

    # Excel 命令
    p = subparsers.add_parser('read-excel', help='读取 Excel 工作簿')
    p.add_argument('-f', '--file', required=True, help='xlsx 文件路径')
    p.add_argument('--format', choices=['markdown', 'text', 'json'], default='markdown')
    p.add_argument('--sheet', help='指定工作表名称')
    p.add_argument('--no-header', action='store_true', help='数据无表头')
    p.add_argument('--header-row', type=int, default=0, help='表头行索引（默认0）')
    p.add_argument('--include-empty', action='store_true', help='包含空行')

    p = subparsers.add_parser('excel-sheets', help='列出 Excel 工作表信息')
    p.add_argument('-f', '--file', required=True, help='xlsx 文件路径')

    p = subparsers.add_parser('excel-table', help='提取指定工作表数据')
    p.add_argument('-f', '--file', required=True, help='xlsx 文件路径')
    p.add_argument('--sheet', help='工作表名称')
    p.add_argument('--no-header', action='store_true', help='数据无表头')
    p.add_argument('--header-row', type=int, default=0, help='表头行索引')

    p = subparsers.add_parser('search-excel', help='搜索 Excel 内容')
    p.add_argument('-f', '--file', required=True, help='xlsx 文件路径')
    p.add_argument('-q', '--query', required=True, help='搜索关键词')
    p.add_argument('--sheet', help='指定工作表名称')
    p.add_argument('--case-sensitive', action='store_true', help='区分大小写')

    args = parser.parse_args()

    if not args.command:
        parser.print_help()
        return

    try:
        handlers = {
            'read-docx': cmd_read_docx,
            'docx-metadata': cmd_docx_metadata,
            'docx-tables': cmd_docx_tables,
            'docx-images': cmd_docx_images,
            'docx-links': cmd_docx_links,
            'docx-charts': cmd_docx_charts,
            'search-docx': cmd_search_docx,
            'read-excel': cmd_read_excel,
            'excel-sheets': cmd_excel_sheets,
            'excel-table': cmd_excel_table,
            'search-excel': cmd_search_excel,
        }
        handlers[args.command](args)
    except FileNotFoundError as e:
        print(f"错误: {e}", file=sys.stderr)
        sys.exit(1)
    except ValueError as e:
        print(f"错误: {e}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"处理文档时发生错误: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()
